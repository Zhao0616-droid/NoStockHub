#!/usr/bin/env python
"""Generate course design report docx for NoStockHub project."""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))
DOCS_DIR = os.path.join(PROJECT_ROOT, 'docs')

PROJECT_NAME = 'NoStockHub（软件项目管理平台）'
TEAM_LEAD = '赵嘉诚'
TEAM_MEMBERS_ALL = [
    ('赵嘉诚', 'U202417380', '组长', '系统设计 + 后端看板/冲刺/文件/报表 + 测试部署'),
    ('王子琪', 'U202417375', '组员', '需求分析 + 前端架构/Layout + 认证与路由'),
    ('路昊天', 'U202417370', '组员', '前端仪表盘 + 项目页 + 公共组件'),
    ('周硕',   'U202417382', '组员', '后端任务 + 工时模块'),
    ('刘经纬', 'U202417369', '组员', '后端认证 + 通知 + 数据库'),
    ('王恒',   'U202417374', '组员', '后端架构 + 项目模块 + 公共基类'),
    ('万晶宇', 'U202417373', '组员', '前端甘特图/看板/任务列表 + 图表组件'),
    ('胡博涵', 'U202415693', '组员', '前端冲刺/报表页 + 部署测试'),
]
TEAM_MEMBER_NAMES = [m[0] for m in TEAM_MEMBERS_ALL]
FILENAME = f'{PROJECT_NAME.split("（")[0]}-{TEAM_LEAD}-{"-".join(TEAM_MEMBER_NAMES[1:])}.docx'

# Report structure: parts → chapters
REPORT_STRUCTURE = [
    ('第一部分：需求分析', [
        ('user_stories.md', '用户故事'),
        ('use_cases.md', '用例交互场景'),
    ]),
    ('第二部分：系统设计', [
        ('architect.md', '架构与类设计'),
        ('db.md', '数据库设计'),
        ('ui_design.md', '前端UI设计'),
        ('backend_api.md', '后端API接口文档'),
    ]),
    ('第三部分：测试与部署', [
        ('test.md', '测试报告'),
        ('install.md', '安装部署文档'),
    ]),
    ('第四部分：使用说明', [
        ('user_guid.md', '软件使用说明书'),
    ]),
    ('第五部分：项目过程记录', [
        ('assign.md', '团队分工记录'),
        ('ai.md', 'AI使用记录'),
    ]),
]

SOURCE_DIRS = [
    ('backend', '后端 Django 代码'),
    ('frontend', '前端 Vue 3 代码'),
    ('sql', '数据库脚本'),
    ('.', '项目根配置文件'),
]

ROOT_ONLY_EXTS = {'.yml', '.yaml', '.example', '.json', '.js', '.html', '.txt'}
ROOT_ONLY_FILES = {'Dockerfile', '.env.example', '.gitignore', 'docker-compose.yml',
                    'docker-compose.prod.yml', 'README.md'}

CODE_EXTS = {'.py', '.js', '.ts', '.vue', '.html', '.css', '.java', '.go',
             '.json', '.yml', '.yaml', '.xml', '.sql', '.sh', '.env', '.md',
             '.txt', '.ini', '.cfg', '.conf'}

SKIP_PATTERNS = [
    'node_modules', '__pycache__', '.git', 'venv', '.idea', '.vscode',
    'package-lock.json', 'migrations', '.claude', 'dist',
]


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), '0')
        tcPr.append(element)


def add_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                           top={'val': 'single', 'sz': '4', 'color': '999999'},
                           bottom={'val': 'single', 'sz': '4', 'color': '999999'},
                           left={'val': 'single', 'sz': '4', 'color': '999999'},
                           right={'val': 'single', 'sz': '4', 'color': '999999'})


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    return heading


def add_part_title(doc, part_title):
    """Add a part separator page."""
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(part_title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.add_paragraph()
    doc.add_page_break()


def parse_markdown_table(lines, start_idx):
    if start_idx + 2 > len(lines):
        return None, start_idx
    header_line = lines[start_idx]
    sep_line = lines[start_idx + 1] if start_idx + 1 < len(lines) else ''
    if '|' not in header_line or '---' not in sep_line:
        return None, start_idx

    headers = [h.strip() for h in header_line.strip().strip('|').split('|')]
    rows = []
    end_idx = start_idx + 2
    while end_idx < len(lines):
        line = lines[end_idx].strip()
        if line.startswith('|') and '|' in line[1:]:
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            rows.append(cells[:len(headers)])
            end_idx += 1
        else:
            break
    return (headers, rows), end_idx


def add_code_block(doc, code_text, language=''):
    for line in code_text.split('\n'):
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(13)
        run = p.add_run(line or ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')


def convert_markdown_to_docx(doc, md_text, base_level=2):
    """Convert markdown text to docx content.
    base_level: docx heading level for markdown `#` headers.
    `#` → base_level, `##` → base_level+1, `###` → base_level+2 (max 3).
    """
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    code_lang = ''

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_buffer), code_lang)
                code_buffer = []
                code_lang = ''
                in_code_block = False
                doc.add_paragraph()
            else:
                in_code_block = True
                code_lang = line.strip()[3:].strip()
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Table
        if '|' in line and line.strip().startswith('|'):
            table_data, end_idx = parse_markdown_table(lines, i)
            if table_data:
                headers, rows = table_data
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                add_table_borders(table)
                for ci, h in enumerate(headers):
                    cell = table.rows[0].cells[ci]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    run = p.add_run(h)
                    run.bold = True
                    run.font.size = Pt(9)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for ri, row in enumerate(rows):
                    for ci, cell_text in enumerate(row):
                        if ci < len(headers):
                            cell = table.rows[ri + 1].cells[ci]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            run = p.add_run(cell_text)
                            run.font.size = Pt(9)
                doc.add_paragraph()
                i = end_idx
                continue

        # Headers: offset by base_level, cap at heading level 3
        h_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if h_match:
            md_level = len(h_match.group(1))
            docx_level = min(base_level + md_level - 1, 3)
            add_heading_styled(doc, h_match.group(2).strip(), docx_level)
            i += 1
            continue

        # Unordered list
        ul_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if ul_match:
            p = doc.add_paragraph(style='List Bullet')
            process_inline_markdown(p, ul_match.group(2))
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r'^(\s*)\d+[\.\)]\s+(.+)$', line)
        if ol_match:
            p = doc.add_paragraph(style='List Number')
            process_inline_markdown(p, ol_match.group(2))
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            doc.add_paragraph('─' * 60)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        process_inline_markdown(p, line)
        i += 1


def process_inline_markdown(paragraph, text):
    pattern = r'\*\*(.+?)\*\*'
    parts = re.split(pattern, text)
    for j, part in enumerate(parts):
        if j % 2 == 1:
            run = paragraph.add_run(part)
            run.bold = True
        else:
            code_pattern = r'`(.+?)`'
            code_parts = re.split(code_pattern, part)
            for k, cp in enumerate(code_parts):
                if k % 2 == 1:
                    run = paragraph.add_run(cp)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
                else:
                    paragraph.add_run(cp)


def add_title_page(doc):
    """Cover page with project info and team table."""
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('软件项目管理平台')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('NoStockHub')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Calibri'

    doc.add_paragraph()
    doc.add_paragraph()

    course = doc.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = course.add_run('软件工程理论与实践课程设计报告')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()
    doc.add_paragraph()

    # Team info table
    rows_count = 4 + len(TEAM_MEMBERS_ALL)  # 3 summary + 1 header + N members
    info_table = doc.add_table(rows=rows_count, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(info_table)

    # Project summary rows
    summary_rows = [
        ('项目名称', '软件项目管理平台（NoStockHub）'),
        ('开发模式', '前后端分离（Vue 3 + Django REST Framework）'),
        ('数据库', 'MySQL 8.0 + Redis 7.0'),
    ]
    for ri, (label, value) in enumerate(summary_rows):
        for ci, text in enumerate([label, value]):
            cell = info_table.rows[ri].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(10)
            if ci == 0:
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Team header row
    team_header = info_table.rows[3]
    for ci, text in enumerate(['角色', '姓名 — 学号 — 负责方向']):
        cell = team_header.cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Team member rows
    for mi, (name, sid, role, duty) in enumerate(TEAM_MEMBERS_ALL):
        row = info_table.rows[4 + mi]
        # Role column
        cell0 = row.cells[0]
        cell0.text = ''
        p0 = cell0.paragraphs[0]
        run0 = p0.add_run(role)
        run0.font.size = Pt(9)
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Name + ID + duty column
        cell1 = row.cells[1]
        cell1.text = ''
        p1 = cell1.paragraphs[0]
        run1 = p1.add_run(f'{name}  {sid}  —  {duty}')
        run1.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('2026年5月')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


def add_toc_page(doc):
    """Add a table of contents page."""
    add_heading_styled(doc, '目  录', level=1)

    chapter_num = 0
    for part_title, chapters in REPORT_STRUCTURE:
        # Part title in TOC
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(part_title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

        for filename, chapter_title in chapters:
            chapter_num += 1
            filepath = os.path.join(DOCS_DIR, filename)
            exists = os.path.exists(filepath)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run(f'第{chapter_num}章  {chapter_title}')
            run.font.size = Pt(10.5)
            if not exists:
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # Appendix in TOC
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('附录：完整源代码')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    doc.add_page_break()


def add_chapter_content(doc, filepath, chapter_num, chapter_title):
    """Add a chapter with proper numbering."""
    add_heading_styled(doc, f'第{chapter_num}章  {chapter_title}', level=1)
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    # Use base_level=2 so doc-internal `#` → h2, `##` → h3, etc.
    convert_markdown_to_docx(doc, content, base_level=2)
    doc.add_page_break()


def collect_source_files():
    result = {}
    for dirpath, dir_label in SOURCE_DIRS:
        full_dir = os.path.join(PROJECT_ROOT, dirpath)
        if not os.path.exists(full_dir):
            continue
        files = []

        if dirpath == '.':
            # Root: only collect specific config files, non-recursive
            for fname in sorted(os.listdir(full_dir)):
                full_path = os.path.join(full_dir, fname)
                if not os.path.isfile(full_path):
                    continue
                ext = os.path.splitext(fname)[1].lower()
                if fname in ROOT_ONLY_FILES or ext in ROOT_ONLY_EXTS:
                    rel_path = os.path.relpath(full_path, PROJECT_ROOT)
                    if os.path.getsize(full_path) > 500 * 1024:
                        continue
                    files.append((rel_path, full_path))
        else:
            for root, dirs, filenames in os.walk(full_dir):
                dirs[:] = [d for d in dirs if not any(p in d for p in SKIP_PATTERNS)]
                for fname in sorted(filenames):
                    ext = os.path.splitext(fname)[1].lower()
                    if ext in CODE_EXTS or fname in ['Dockerfile', '.env.example', '.gitignore']:
                        full_path = os.path.join(root, fname)
                        rel_path = os.path.relpath(full_path, PROJECT_ROOT)
                        if os.path.getsize(full_path) > 500 * 1024:
                            continue
                        files.append((rel_path, full_path))

        result[dir_label] = sorted(files, key=lambda x: x[0])
    return result


def add_source_appendix(doc):
    doc.add_page_break()
    add_heading_styled(doc, '附录：完整源代码', level=1)

    source_files = collect_source_files()

    for dir_label, files in source_files.items():
        add_heading_styled(doc, dir_label, level=2)

        for rel_path, full_path in files:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f'// {rel_path}')
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

            try:
                with open(full_path, 'r', encoding='utf-8') as f:
                    code = f.read()
                if len(code) > 10000:
                    code = code[:10000] + '\n// ... (truncated)'
                add_code_block(doc, code, os.path.splitext(full_path)[1].lower().lstrip('.'))
            except Exception:
                p = doc.add_paragraph()
                run = p.add_run('[无法读取文件]')
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

            doc.add_paragraph()


def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # Default font
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 1. Title page
    add_title_page(doc)

    # 2. Table of contents
    add_toc_page(doc)

    # 3. Chapters organized by parts
    chapter_num = 0
    for part_title, chapters in REPORT_STRUCTURE:
        add_part_title(doc, part_title)
        for filename, chapter_title in chapters:
            chapter_num += 1
            filepath = os.path.join(DOCS_DIR, filename)
            if os.path.exists(filepath):
                print(f'Adding: 第{chapter_num}章 {chapter_title}  ({filename})')
                add_chapter_content(doc, filepath, chapter_num, chapter_title)
            else:
                print(f'Skipping (not found): {filename}')

    # 4. Source code appendix
    print('Adding source code appendix...')
    add_source_appendix(doc)

    # Save
    output_path = os.path.join(PROJECT_ROOT, FILENAME)
    doc.save(output_path)
    print(f'\nReport generated: {output_path}')
    print(f'File size: {os.path.getsize(output_path) / 1024 / 1024:.1f} MB')


if __name__ == '__main__':
    main()
